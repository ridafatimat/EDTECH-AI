from __future__ import annotations

import argparse
import json
import re
import sys
import time
from dataclasses import asdict, is_dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from dotenv import load_dotenv

from app.db.repositories.technical_correction_repository import (
    PostgreSQLTechnicalCorrectionRepository,
)
from app.db.session import session_scope
from app.services.semantic_chunker import SemanticChunker, SemanticChunkingConfig
from app.services.selective_technical_normalizer import SelectiveTechnicalNormaliser
from app.services.technical_correction_client import GroqTechnicalCorrectionClient
from app.services.topic_pipeline import Module3TopicPipeline
from app.services.transcript_preprocessor import preprocess_transcript


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT_ROOT = PROJECT_ROOT / "test_outputs" / "agent1_two_transcript_docx_test"


# -----------------------------------------------------------------------------
# Generic helpers
# -----------------------------------------------------------------------------

def to_plain_data(value: Any) -> Any:
    if value is None:
        return None

    if isinstance(value, Path):
        return str(value)

    if isinstance(value, datetime):
        return value.isoformat()

    if hasattr(value, "model_dump"):
        return to_plain_data(value.model_dump())

    if hasattr(value, "dict"):
        try:
            return to_plain_data(value.dict())
        except TypeError:
            pass

    if hasattr(value, "to_dict"):
        return to_plain_data(value.to_dict())

    if is_dataclass(value):
        return to_plain_data(asdict(value))

    if isinstance(value, dict):
        return {
            str(key): to_plain_data(item)
            for key, item in value.items()
        }

    if isinstance(value, (list, tuple, set)):
        return [to_plain_data(item) for item in value]

    if isinstance(value, (str, int, float, bool)):
        return value

    if hasattr(value, "__dict__"):
        return {
            key: to_plain_data(item)
            for key, item in vars(value).items()
            if not key.startswith("_")
        }

    return str(value)


def safe_name(path: Path) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_-]+", "_", path.stem.strip())
    cleaned = re.sub(r"_+", "_", cleaned).strip("_")
    return cleaned or "transcript"


def resolve_project_path(path: Path) -> Path:
    return path if path.is_absolute() else PROJECT_ROOT / path


def extract_docx_text(docx_path: Path) -> str:
    if not docx_path.exists():
        raise FileNotFoundError(f"Input DOCX not found: {docx_path}")

    if docx_path.suffix.casefold() != ".docx":
        raise ValueError(f"Expected a DOCX file: {docx_path}")

    document = Document(docx_path)
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    text = "\n".join(paragraphs).strip()

    if not text:
        raise ValueError(f"No text found inside: {docx_path}")

    return text


def save_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(to_plain_data(payload), indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


def set_default_doc_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)


def add_metadata_table(document: Document, rows: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = "" if value is None else str(value)


def add_text_preserving_lines(document: Document, text: str) -> None:
    for line in text.splitlines():
        document.add_paragraph(line)


# -----------------------------------------------------------------------------
# DOCX outputs
# -----------------------------------------------------------------------------

def save_cleaned_docx(
    *,
    output_path: Path,
    source_file: Path,
    raw_text: str,
    deterministic_text: str,
    final_cleaned_text: str,
    preprocessing_payload: dict[str, Any],
    technical_payload: dict[str, Any],
) -> None:
    document = Document()
    set_default_doc_style(document)

    document.add_heading("Agent 1 — Module 1 Cleaning Result", level=0)
    add_metadata_table(
        document,
        [
            ("Source transcript", source_file.name),
            ("Original characters", len(raw_text)),
            ("Deterministically cleaned characters", len(deterministic_text)),
            ("Final cleaned characters", len(final_cleaned_text)),
            ("Generated at", datetime.now().astimezone().isoformat()),
        ],
    )

    document.add_heading("Final Cleaned Transcript", level=1)
    add_text_preserving_lines(document, final_cleaned_text)

    document.add_page_break()
    document.add_heading("Deterministic Preprocessing Audit", level=1)
    document.add_paragraph(
        json.dumps(preprocessing_payload, indent=2, ensure_ascii=False)
    )

    document.add_heading("Technical Normalisation Audit", level=1)
    document.add_paragraph(
        json.dumps(technical_payload, indent=2, ensure_ascii=False)
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def save_chunks_docx(
    *,
    output_path: Path,
    source_file: Path,
    chunking_payload: dict[str, Any],
) -> None:
    document = Document()
    set_default_doc_style(document)

    chunks = chunking_payload.get("chunks", [])

    document.add_heading("Agent 1 — Module 2 Semantic Chunking Result", level=0)
    add_metadata_table(
        document,
        [
            ("Source transcript", source_file.name),
            ("Embedding model", chunking_payload.get("embedding_model")),
            ("Semantic threshold", chunking_payload.get("semantic_threshold")),
            ("Total sentences", chunking_payload.get("total_sentences")),
            ("Total words", chunking_payload.get("total_words")),
            ("Final chunk count", len(chunks)),
        ],
    )

    if not chunks:
        document.add_paragraph("No chunks were produced.")
    else:
        for index, chunk in enumerate(chunks):
            document.add_heading(
                f"Chunk {chunk.get('chunk_id', index + 1)}",
                level=1,
            )

            add_metadata_table(
                document,
                [
                    ("Word count", chunk.get("word_count")),
                    ("Sentence count", chunk.get("sentence_count")),
                    (
                        "Sentence range",
                        f"{chunk.get('start_sentence')} → {chunk.get('end_sentence')}",
                    ),
                    ("Boundary reason", chunk.get("boundary_reason")),
                    ("Boundary similarity", chunk.get("boundary_similarity")),
                    ("Overlap words", chunk.get("overlap_word_count", 0)),
                ],
            )

            document.add_paragraph(str(chunk.get("text", "")))

            if index < len(chunks) - 1:
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def save_topics_docx(
    *,
    output_path: Path,
    source_file: Path,
    topic_payload: dict[str, Any],
) -> None:
    document = Document()
    set_default_doc_style(document)

    merged_topics = topic_payload.get("merged_topics", [])
    chunk_results = topic_payload.get("chunk_results", [])

    document.add_heading("Agent 1 — Module 3 Topic Mapping Result", level=0)
    add_metadata_table(
        document,
        [
            ("Source transcript", source_file.name),
            ("Embedding model", topic_payload.get("embedding_model")),
            ("Total chunks", topic_payload.get("total_chunks")),
            ("CS-relevant chunks", topic_payload.get("cs_relevant_chunks")),
            ("Non-CS/no-new-topic chunks", topic_payload.get("non_cs_chunks")),
            ("Merged topic count", len(merged_topics)),
        ],
    )

    document.add_heading("Merged AQA Lesson Topics", level=1)

    if not merged_topics:
        document.add_paragraph(
            "No AQA Computer Science topics were retained from this transcript."
        )
    else:
        for number, topic in enumerate(merged_topics, start=1):
            document.add_heading(
                f"{number}. {topic.get('topic', 'Unnamed topic')}",
                level=2,
            )

            add_metadata_table(
                document,
                [
                    ("Concept ID", topic.get("concept_id")),
                    ("Domain", topic.get("domain")),
                    ("Topic role", topic.get("topic_role")),
                    ("Confidence", topic.get("confidence")),
                    ("Ranking score", topic.get("ranking_score")),
                    ("Source chunk IDs", topic.get("source_chunk_ids")),
                    ("Syllabus topic", topic.get("syllabus_topic")),
                    ("Syllabus subtopic", topic.get("syllabus_subtopic")),
                    ("Specification reference", topic.get("specification_reference")),
                ],
            )

            evidence = topic.get("evidence") or topic.get("supporting_text")
            if evidence:
                document.add_paragraph(f"Evidence: {evidence}")

    document.add_page_break()
    document.add_heading("Chunk-by-Chunk Classification", level=1)

    if not chunk_results:
        document.add_paragraph("No chunk-level topic results were returned.")
    else:
        for result in chunk_results:
            document.add_heading(
                f"Chunk {result.get('chunk_id')}",
                level=2,
            )

            add_metadata_table(
                document,
                [
                    ("Classification", result.get("classification")),
                    ("CS relevant", result.get("is_cs_relevant")),
                    ("Creates new topic", result.get("creates_new_topic")),
                    ("Requires LLM fallback", result.get("requires_llm_fallback")),
                ],
            )

            candidates = result.get("topic_candidates", [])
            if not candidates:
                document.add_paragraph("Retained topics: none")
            else:
                for candidate in candidates:
                    document.add_paragraph(
                        (
                            f"• {candidate.get('topic')} | "
                            f"confidence={candidate.get('confidence')} | "
                            f"concept_id={candidate.get('concept_id')}"
                        )
                    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


# -----------------------------------------------------------------------------
# Pipeline execution
# -----------------------------------------------------------------------------

def run_single_transcript(
    *,
    source_file: Path,
    output_root: Path,
    no_llm: bool,
) -> dict[str, Any]:
    run_name = safe_name(source_file)
    run_dir = output_root / run_name
    run_dir.mkdir(parents=True, exist_ok=True)

    cleaned_docx = run_dir / "01_cleaned_transcript.docx"
    chunks_docx = run_dir / "02_semantic_chunks.docx"
    topics_docx = run_dir / "03_topic_mapping.docx"
    raw_json = run_dir / "pipeline_raw_outputs.json"

    started = time.perf_counter()

    # MODULE 1A: DOCX extraction + deterministic cleaning
    raw_text = extract_docx_text(source_file)
    preprocessing_result = preprocess_transcript(raw_text=raw_text)
    deterministic_text = preprocessing_result.cleaned_text.strip()

    if not deterministic_text:
        raise ValueError("Deterministic preprocessing returned empty text.")

    # MODULE 1B: technical terminology normalisation
    correction_client = None if no_llm else GroqTechnicalCorrectionClient()

    with session_scope() as session:
        repository = PostgreSQLTechnicalCorrectionRepository(session)
        normaliser = SelectiveTechnicalNormaliser(
            repository=repository,
            correction_client=correction_client,
        )
        technical_result = normaliser.normalise(deterministic_text)

    final_cleaned_text = technical_result.normalised_text.strip()

    if not final_cleaned_text:
        raise ValueError("Technical normalisation returned empty text.")

    preprocessing_payload = to_plain_data(preprocessing_result)
    technical_payload = to_plain_data(technical_result)

    save_cleaned_docx(
        output_path=cleaned_docx,
        source_file=source_file,
        raw_text=raw_text,
        deterministic_text=deterministic_text,
        final_cleaned_text=final_cleaned_text,
        preprocessing_payload=preprocessing_payload,
        technical_payload=technical_payload,
    )

    # MODULE 2: semantic chunking
    config = SemanticChunkingConfig(
        min_chunk_words=150,
        target_chunk_words=325,
        max_chunk_words=550,
        strong_transition_min_words=80,
        semantic_unit_words=60,
        boundary_percentile=15.0,
        threshold_floor=0.10,
        threshold_ceiling=0.45,
        soft_transition_margin=0.10,
        soft_transition_similarity_ceiling=0.35,
        max_size_overlap_words=45,
        max_size_overlap_sentences=2,
    )

    chunker = SemanticChunker(config=config)
    chunking_result = chunker.chunk(final_cleaned_text)
    chunking_payload = to_plain_data(chunking_result)

    save_chunks_docx(
        output_path=chunks_docx,
        source_file=source_file,
        chunking_payload=chunking_payload,
    )

    # MODULE 3: topic extraction and syllabus mapping
    saved_chunks = chunking_payload.get("chunks", [])
    topic_pipeline = Module3TopicPipeline()
    topic_result = topic_pipeline.process_chunks(saved_chunks)
    topic_payload = to_plain_data(topic_result)

    save_topics_docx(
        output_path=topics_docx,
        source_file=source_file,
        topic_payload=topic_payload,
    )

    result = {
        "source_file": str(source_file.resolve()),
        "status": "completed",
        "duration_seconds": round(time.perf_counter() - started, 3),
        "cleaned_docx": str(cleaned_docx.resolve()),
        "chunks_docx": str(chunks_docx.resolve()),
        "topics_docx": str(topics_docx.resolve()),
        "chunk_count": len(chunking_payload.get("chunks", [])),
        "cs_relevant_chunks": topic_payload.get("cs_relevant_chunks"),
        "non_cs_chunks": topic_payload.get("non_cs_chunks"),
        "merged_topic_count": len(topic_payload.get("merged_topics", [])),
        "raw_outputs": {
            "preprocessing": preprocessing_payload,
            "technical_normalisation": technical_payload,
            "chunking": chunking_payload,
            "topics": topic_payload,
        },
    }

    save_json(raw_json, result)
    return result


def save_batch_summary_docx(output_path: Path, results: list[dict[str, Any]]) -> None:
    document = Document()
    set_default_doc_style(document)

    document.add_heading("Agent 1 — Two Transcript Test Summary", level=0)
    document.add_paragraph(
        "This batch checks one AQA Computer Science transcript and one transcript "
        "containing no Computer Science lesson content."
    )

    for result in results:
        document.add_heading(Path(result["source_file"]).name, level=1)
        add_metadata_table(
            document,
            [
                ("Status", result.get("status")),
                ("Duration (seconds)", result.get("duration_seconds")),
                ("Chunk count", result.get("chunk_count")),
                ("CS-relevant chunks", result.get("cs_relevant_chunks")),
                ("Non-CS chunks", result.get("non_cs_chunks")),
                ("Merged topics", result.get("merged_topic_count")),
                ("Cleaned DOCX", result.get("cleaned_docx")),
                ("Chunks DOCX", result.get("chunks_docx")),
                ("Topics DOCX", result.get("topics_docx")),
            ],
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    load_dotenv()

    parser = argparse.ArgumentParser(
        description=(
            "Test two Agent 1 transcript DOCX files and generate cleaned, "
            "chunked and topic-mapping DOCX outputs for each file."
        )
    )

    parser.add_argument(
        "--files",
        nargs="+",
        required=True,
        type=Path,
        help="Paths to the two transcript DOCX files.",
    )

    parser.add_argument(
        "--output-root",
        type=Path,
        default=DEFAULT_OUTPUT_ROOT,
        help="Folder in which test outputs will be stored.",
    )

    parser.add_argument(
        "--no-llm",
        action="store_true",
        help=(
            "Disable Groq technical correction. PostgreSQL correction memory "
            "and deterministic normalisation still run."
        ),
    )

    args = parser.parse_args()

    if len(args.files) != 2:
        parser.error("Exactly two DOCX files must be supplied to --files.")

    source_files = [resolve_project_path(path) for path in args.files]
    output_root = resolve_project_path(args.output_root)
    output_root.mkdir(parents=True, exist_ok=True)

    results: list[dict[str, Any]] = []
    failures: list[dict[str, str]] = []

    print("=" * 100)
    print("AGENT 1 — TWO TRANSCRIPT DOCX TEST")
    print("=" * 100)

    for source_file in source_files:
        print(f"\nProcessing: {source_file}")

        try:
            result = run_single_transcript(
                source_file=source_file,
                output_root=output_root,
                no_llm=args.no_llm,
            )
            results.append(result)
            print(f"Completed: {source_file.name}")
            print(f"  Cleaned: {result['cleaned_docx']}")
            print(f"  Chunks:  {result['chunks_docx']}")
            print(f"  Topics:  {result['topics_docx']}")
        except Exception as exc:
            failures.append(
                {
                    "source_file": str(source_file),
                    "error_type": type(exc).__name__,
                    "message": str(exc),
                }
            )
            print(f"FAILED: {source_file.name}")
            print(f"  {type(exc).__name__}: {exc}")

    summary_json = output_root / "batch_summary.json"
    summary_docx = output_root / "batch_summary.docx"

    save_json(
        summary_json,
        {
            "completed": results,
            "failed": failures,
        },
    )

    if results:
        save_batch_summary_docx(summary_docx, results)

    print("\n" + "=" * 100)
    print("TEST FINISHED")
    print("=" * 100)
    print(f"Output root: {output_root.resolve()}")
    print(f"Completed files: {len(results)}")
    print(f"Failed files: {len(failures)}")
    print(f"Summary JSON: {summary_json.resolve()}")
    if results:
        print(f"Summary DOCX: {summary_docx.resolve()}")
    print("=" * 100)

    if failures:
        sys.exit(1)


if __name__ == "__main__":
    main()